"""
Сборка итогового отчёта ЛР-3 в формате DOCX:
- титульный лист «Лабораторная работа - 3»;
- раздел ЗАДАНИЕ с формулами P_ij ≈ e^(-a*d^b) и P_ij ≈ 1/d^b;
- в конце — полный листинг программы (lab-3.py).
Требуется: pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images")
INTERESTING_DIR = os.path.join(IMAGES_DIR, "interesting_4x5")
LAB_PY = os.path.join(SCRIPT_DIR, "lab-3.py")
OUT_DOCX = os.path.join(SCRIPT_DIR, "lab-3-otchet.docx")


def set_document_styles(doc):
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"
    pfr = style.paragraph_format
    pfr.space_after = Pt(6)
    pfr.first_line_indent = Cm(1.25)
    pfr.line_spacing = 1.5
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(14 if level == 1 else 12)
        h.paragraph_format.first_line_indent = Cm(0)


def paragraph_no_indent(doc, text, align=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0)
    if align is not None:
        p.alignment = align
    return p


def add_title_page(doc):
    def center_para(text, bold=False, size=Pt(12)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.size = size
        r.font.bold = bold
        r.font.name = "Times New Roman"

    doc.add_paragraph()
    center_para(
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ\n"
        "ВЫСШЕГО ОБРАЗОВАНИЯ «СИБИРСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ\n"
        "ТЕЛЕКОММУНИКАЦИЙ И ИНФОРМАТИКИ»",
        bold=True,
    )
    center_para("Лабораторная работа - 3\nпо дисциплине «Моделирование»", size=Pt(14))
    center_para("Генерация графов, похожих на реальные сети")
    center_para("Выполнил студент\nНиколаенков М. Д.\nГруппы ИВ-222")
    center_para("Новосибирск – 2025")


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    set_document_styles(doc)
    add_title_page(doc)
    doc.add_page_break()

    doc.add_heading("ЗАДАНИЕ", level=1)
    paragraph_no_indent(
        doc,
        "Написать код, генерирующий случайные точки на плоскости 100×100 и строящий на этих точках граф. "
        "Для построения ребра используются две формулы вероятности (d_ij — расстояние между точками i и j; циклы недопустимы; выбор вершины, из которой строится ребро, рандомизируется; вероятности пересчитываются каждый раз):"
    )
    paragraph_no_indent(doc, "1) P_ij ≈ e^(-a*d^b); параметры a и b варьируются.")
    paragraph_no_indent(doc, "2) P_ij ≈ 1/d^b; варьируется параметр b.")
    paragraph_no_indent(doc, "Сгенерировать 4 набора конфигураций (например, 3 комбинации a и b для первой формулы и 1 конфигурация с параметром b для второй), по 5 графов на каждый набор — всего 20 графов.")

    doc.add_heading("Выполнение", level=1)
    paragraph_no_indent(
        doc,
        "Реализованы обе формулы. Для первой (экспоненциальной) варьируются a и b; для второй (степенной) — только b. "
        "Графы строятся без циклов (объединение компонент через UnionFind), вероятности нормализуются при выборе ребра."
    )
    sample_img = os.path.join(INTERESTING_DIR, "1_v1_exp_global", "graph_1.png")
    if os.path.isfile(sample_img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run()
        r.add_picture(sample_img, width=Cm(12))
        cap = doc.add_paragraph("Рисунок 1 – Пример графа (вариация 1: exp(-a*d^b), a=0.002, b=0.6)")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
    paragraph_no_indent(doc, "Все 20 графов (4 вариации × 5 штук) сохранены в каталоге images/interesting_4x5/.")

    doc.add_page_break()
    doc.add_heading("ЛИСТИНГ ПРОГРАММЫ", level=1)
    if os.path.isfile(LAB_PY):
        with open(LAB_PY, "r", encoding="utf-8") as f:
            code_text = f.read()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(code_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    else:
        paragraph_no_indent(doc, f"[Файл не найден: {LAB_PY}]")

    doc.save(OUT_DOCX)
    print(f"Отчёт сохранён: {OUT_DOCX}")


if __name__ == "__main__":
    main()
