"""
Генерация отчёта лабы 2 в формате DOCX: титульник, два графа, расчёт CPM.
Требуется: pip install python-docx
"""
import os
import sys
import importlib.util
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images")
OUT_DOCX = os.path.join(SCRIPT_DIR, "lab-2-otchet.docx")

# Единый источник: графы и расчёт CPM из lab-2.py
spec = importlib.util.spec_from_file_location("lab2", os.path.join(SCRIPT_DIR, "lab-2.py"))
lab2 = importlib.util.module_from_spec(spec)
sys.modules["lab2"] = lab2
spec.loader.exec_module(lab2)
GRAPH_VARIANT_17 = lab2.GRAPH_VARIANT_17
GRAPH_CUSTOM = lab2.GRAPH_CUSTOM
run_cpm = lab2.run_cpm


def set_document_styles(doc):
    """Форматирование по ГОСТ: шрифт, интервал 1,5, абзацный отступ."""
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"
    pfr = style.paragraph_format
    pfr.space_after = Pt(6)
    pfr.first_line_indent = Cm(1.25)
    pfr.line_spacing = 1.5  # полуторный интервал по ГОСТ
    for heading_level in range(1, 4):
        h = doc.styles[f"Heading {heading_level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(14 if heading_level == 1 else 12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.first_line_indent = Cm(0)


def paragraph_no_indent(doc, text, align=None):
    """Абзац без красной строки (для списков и подписей)."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0)
    if align is not None:
        p.alignment = align
    return p


def add_title_page(doc):
    """Титульная страница (по образцу ЛР-1)."""
    def center_para(text, bold=False, size=Pt(12)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.size = size
        r.font.bold = bold
        r.font.name = "Times New Roman"
    doc.add_paragraph()
    center_para(
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ\n"
        "ВЫСШЕГО ОБРАЗОВАНИЯ «СИБИРСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ\n"
        "ТЕЛЕКОММУНИКАЦИЙ И ИНФОРМАТИКИ»",
        bold=True,
    )
    center_para("Лабораторная работа № 2\nпо дисциплине «Моделирование»\nВариант — 17", size=Pt(14))
    center_para("Выполнил студент\nНиколаенков М. Д.\nГруппы ИВ-222")
    center_para("Новосибирск – 2025")


def add_heading(doc, text, level=1):
    doc.add_paragraph()
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)


def add_graph_section(doc, jobs, figure_num, table_num, title, image_filename, description, steps_lines, figure_caption, table_caption):
    """Секция по одному графу: описание, Рисунок N – ..., пошаговый расчёт, Таблица N – ..., итог. Возвращает (figure_num+1, table_num+1)."""
    add_heading(doc, title, level=2)
    paragraph_no_indent(doc, description)

    image_path = os.path.join(IMAGES_DIR, image_filename)
    if os.path.isfile(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run()
        r.add_picture(image_path, width=Inches(5.2))
        cap = doc.add_paragraph(f"Рисунок {figure_num} – {figure_caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
    else:
        paragraph_no_indent(doc, f"[Рисунок не найден: {image_filename}. Запустите draw_graphs.py.]")

    add_heading(doc, "Расчёт по шагам (метод CPM)", level=3)
    paragraph_no_indent(doc, "Расчёт выполнен по стандартному алгоритму: прямой ход (ранние сроки событий), обратный ход (поздние сроки), затем для каждой работы — раннее/позднее начало и окончание, резервы; критический путь определяется по работам с нулевым полным резервом.")
    res = run_cpm(jobs)

    # Шифр буквами: узел 1→A, 2→B, … (A–B, A–C, B–D, …)
    def node_letter(i):
        return chr(64 + i)

    def cipher(u, v):
        return f"{node_letter(u)}–{node_letter(v)}"

    # Шаг 4: полное перечисление работ с R > 0 и r > 0 (в обозначениях Шифр)
    r_pos = []
    r_priv_pos = []
    for name, u, v, t, rn, ro, pn, po, R, r_priv in res["rows"]:
        c = cipher(u, v)
        if R > 0:
            r_pos.append((c, R))
        if r_priv > 0:
            r_priv_pos.append((c, r_priv))
    step4_r = ", ".join(f"{c} (R={R})" for c, R in r_pos) if r_pos else "нет"
    step4_r_priv = ", ".join(f"{c} (r={r})" for c, r in r_priv_pos) if r_priv_pos else "нет"
    step4_text = (
        "Шаг 4. Резервы времени. Полный резерв R = ПО − РО (на сколько можно сдвинуть работу, не увеличивая срок проекта). "
        "Частный резерв r = РС(j) − РО (на сколько можно сдвинуть окончание, не сдвигая ранние сроки следующих работ). "
        f"Работы с R = 0 образуют критический путь. Работы с R > 0: {step4_r}. Работы с r > 0: {step4_r_priv}."
    )
    for s in steps_lines:
        if s is None:
            paragraph_no_indent(doc, step4_text)
        elif s:
            paragraph_no_indent(doc, s)
    add_heading(doc, "Итоговая таблица работ", level=3)
    rows = res["rows"]
    table = doc.add_table(rows=len(rows) + 1, cols=8)
    table.style = "Table Grid"
    headers = ["Шифр", "tij", "РН", "РО", "ПН", "ПО", "Rij", "rij"]
    for col, h in enumerate(headers):
        c = table.rows[0].cells[col]
        c.text = h
        for p in c.paragraphs:
            p.paragraph_format.first_line_indent = Cm(0)
    for i, (name, u, v, t, rn, ro, pn, po, R_full, r_priv) in enumerate(rows):
        row = table.rows[i + 1]
        row.cells[0].text = cipher(u, v)
        row.cells[1].text = str(t)
        row.cells[2].text = str(rn)
        row.cells[3].text = str(ro)
        row.cells[4].text = str(pn)
        row.cells[5].text = str(po)
        row.cells[6].text = str(R_full)
        row.cells[7].text = str(r_priv)
    doc.add_paragraph()
    tab_cap = doc.add_paragraph(f"Таблица {table_num} – {table_caption}")
    tab_cap.paragraph_format.first_line_indent = Cm(0)
    path_letters = " → ".join(node_letter(n) for n in res["critical_nodes"])
    paragraph_no_indent(doc, f"Критическое время проекта: {res['total_time']}. Критический путь: {path_letters}.")
    return figure_num + 1, table_num + 1


def main():
    doc = Document()
    # Поля по ГОСТ: левое 3 см, правое 1,5 см, верхнее и нижнее 2 см
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    set_document_styles(doc)
    add_title_page(doc)
    doc.add_page_break()

    add_heading(doc, "Цель работы", level=1)
    paragraph_no_indent(
        doc,
        "Освоить метод критического пути (CPM) для расчёта сетевых графиков: научиться определять ранние и поздние сроки событий, "
        "резервы времени работ и критический путь проекта на примере двух графов (по варианту и собственного)."
    )

    add_heading(doc, "Задание", level=1)
    paragraph_no_indent(doc, "1. Построить и рассчитать методом CPM два сетевых графика: по варианту 17 и собственный (повышенной сложности).")
    paragraph_no_indent(doc, "2. Для каждого графа описать структуру, построить визуализацию (узлы — буквенные обозначения, на рёбрах — веса), выполнить пошаговый расчёт сроков и резервов, указать критический путь.")
    paragraph_no_indent(doc, "3. Оформить отчёт по ГОСТ с титульным листом, рисунками и таблицами.")

    add_heading(doc, "Краткие теоретические сведения", level=1)
    paragraph_no_indent(
        doc,
        "Метод критического пути (CPM) применяется к сетевым графикам проектов. События нумеруются; работы задаются парами (начало, конец) и длительностями."
    )
    paragraph_no_indent(
        doc,
        "Прямой ход: ранний срок события j равен максимуму из (ранний срок предшественника + длительность работы). Обратный ход: поздний срок события i — минимум из (поздний срок преемника − длительность работы). Критическое время Tкр — ранний срок конечного события."
    )
    paragraph_no_indent(
        doc,
        "Для каждой работы: РН — раннее начало, РО — раннее окончание, ПН — позднее начало, ПО — позднее окончание. Полный резерв R = ПО − РО; частный резерв r = РС(j) − РО. Критический путь составляют работы с R = 0."
    )

    add_heading(doc, "Выполнение работы", level=1)

    # Граф по варианту 17 (одобренная таблица: A–B(5), A–C(3), B–D(6), C–E(4), D–F(5), E–F(5); Tкр=16, путь A→B→D→F)
    desc17 = (
        "Схема графа по варианту 17: шесть узлов A–F. Работы: A–B(5), A–C(3), B–D(6), C–E(4), D–F(5), E–F(5). "
        "Две ветви из A, два входа в F. Критический путь на схеме выделен красным: A → B → D → F."
    )
    steps17 = [
        "Шаг 1. Прямой ход — ранние сроки событий. РС(1) = 0. РС(2): A–B, t = 5, РС(2) = 5. РС(3): A–C, t = 3, РС(3) = 3. РС(4): B–D, t = 6, РС(4) = 5 + 6 = 11. РС(5): C–E, t = 4, РС(5) = 3 + 4 = 7. РС(6): входят D–F и E–F; от 4: 11 + 5 = 16, от 5: 7 + 5 = 12; РС(6) = 16. Tкр = 16.",
        "Шаг 2. Обратный ход — поздние сроки. ПС(6) = 16. ПС(5): E–F, ПС(5) = 16 − 5 = 11. ПС(4): D–F, ПС(4) = 16 − 5 = 11. ПС(3): C–E, ПС(3) = 11 − 4 = 7. ПС(2): B–D, ПС(2) = 11 − 6 = 5. ПС(1): выходят A–B и A–C; минимум ПС(1) = 0.",
        "Шаг 3. Для каждой работы: РН = РС(i), РО = РН + t, ПО = ПС(j), ПН = ПО − t. Все значения в итоговой таблице ниже.",
        None,  # Шаг 4 генерируется из res в add_graph_section
        "Шаг 5. Критический путь — работы с R = 0: A–B, B–D, D–F; по узлам: A → B → D → F. Tкр = 16.",
    ]
    figure_num, table_num = add_graph_section(
        doc, GRAPH_VARIANT_17, 1, 1,
        "Граф по варианту 17",
        "graph_variant17.png",
        desc17,
        steps17,
        "Схема графа по варианту 17 (узлы A–F, веса на рёбрах)",
        "Итоговые параметры работ для графа по варианту 17",
    )

    doc.add_page_break()

    # Собственный граф по фото: A,B — старты; C(A,4), D(B,1), E(B,5), F(D,C,2), G(E,F,3), H(G,4); Tкр=13, путь A→C→F→G→H
    desc_custom = (
        "Собственный граф: A, B без предшественника; C(A,4), D(B,1), E(B,5), F(D,C,2), G(E,F,3), H(G,4). "
        "Критический путь на схеме выделен красным: A → C → F → G → H."
    )
    steps_custom = [
        "Шаг 1. Прямой ход. РС(1) = 0, РС(2) = 0. РС(3): A–C, 4. РС(4): B–D, 1. РС(5): B–E, 5. РС(6): C–F и D–F; max(4+2, 1+2) = 6. РС(7): E–G и F–G; max(5+3, 6+3) = 9. РС(8): G–H, 9+4 = 13. Tкр = 13.",
        "Шаг 2. Обратный ход. ПС(8) = 13, ПС(7) = 9, ПС(6) = 6, ПС(5) = 6, ПС(4) = 4, ПС(3) = 4, ПС(2) = 0, ПС(1) = 0.",
        "Шаг 3. Для каждой работы: РН = РС(i), РО = РН + t, ПО = ПС(j), ПН = ПО − t. Результаты в итоговой таблице ниже.",
        None,  # Шаг 4 генерируется из res в add_graph_section
        "Шаг 5. Критический путь — работы с R = 0: A–C, C–F, F–G, G–H; по узлам: A → C → F → G → H. Tкр = 13.",
    ]
    add_graph_section(
        doc, GRAPH_CUSTOM, figure_num, table_num,
        "Собственный граф (повышенная сложность)",
        "graph_custom.png",
        desc_custom,
        steps_custom,
        "Собственный граф (узлы A–H, веса на рёбрах; критический путь выделен красным)",
        "Итоговые параметры работ для собственного графа",
    )

    doc.add_page_break()
    add_heading(doc, "Выводы", level=1)
    res17 = run_cpm(GRAPH_VARIANT_17)
    res_custom = run_cpm(GRAPH_CUSTOM)
    def node_letter(i):
        return chr(64 + i)
    path_17 = " → ".join(node_letter(n) for n in res17["critical_nodes"])
    path_custom = " → ".join(node_letter(n) for n in res_custom["critical_nodes"])
    paragraph_no_indent(doc, "1. Для графа по варианту 17 получено критическое время проекта Tкр = " + str(res17["total_time"]) + "; критический путь: " + path_17 + ".")
    paragraph_no_indent(doc, "2. Для собственного графа (8 работ, 7 событий) Tкр = " + str(res_custom["total_time"]) + "; критический путь: " + path_custom + ".")
    paragraph_no_indent(doc, "3. Визуализации выполнены в едином стиле: узлы — буквенные обозначения, на рёбрах — веса; на втором графе критический путь выделен красным.")
    paragraph_no_indent(doc, "4. Метод CPM позволил однозначно определить ранние и поздние сроки событий, полный и частный резервы работ и длительность проекта для обоих графов.")
    paragraph_no_indent(doc, "5. Отчёт оформлен в соответствии с требованиями ГОСТ (поля, шрифт, интервал, нумерация рисунков и таблиц) и пригоден для презентации и защиты.")

    doc.save(OUT_DOCX)
    print(f"Отчёт сохранён: {OUT_DOCX}")


if __name__ == "__main__":
    main()
